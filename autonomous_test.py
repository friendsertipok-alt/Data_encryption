import os
import httpx
import asyncio
from docx import Document
from openpyxl import Workbook
import fitz  # PyMuPDF

os.makedirs("test_files", exist_ok=True)
os.makedirs("test_results", exist_ok=True)

# 1. Generate TXT (Complex Russian declensions + mixed English)
with open("test_files/test.txt", "w", encoding="utf-8") as f:
    f.write("Настоящий договор заключен с Ивановым Иваном Ивановичем. "
            "Исполнитель: Сидорову Сидору Сидоровичу. "
            "Также в копии John Smith (john.smith@fbi.gov).")

# 2. Generate DOCX
doc = Document()
doc.add_paragraph("Конфиденциальный акт.")
doc.add_paragraph("Сдал: Петровым Петром Петровичем. Принял: Алексееву Алексею Алексеевичу.")
doc.save("test_files/test.docx")

# 3. Generate XLSX
wb = Workbook()
ws = wb.active
ws["A1"] = "Имя"
ws["B1"] = "Email"
ws["A2"] = "Смирнову Смирну Смирновичу"
ws["B2"] = "smirnov@mail.ru"
wb.save("test_files/test.xlsx")

# 4. Generate PDF
pdf_doc = fitz.open()
page = pdf_doc.new_page()
page.insert_text(fitz.Point(50, 72), "Top secret: agent Michael Jordan, phone +1-800-555-0199", fontsize=11)
pdf_doc.save("test_files/test.pdf")

async def test_file(filename):
    url_anon = "http://127.0.0.1:8001/api/anonymize/file"
    url_deanon = "http://127.0.0.1:8001/api/deanonymize/file"
    file_path = f"test_files/{filename}"
    
    with open(file_path, "rb") as f:
        files = {"file": (filename, f, "application/octet-stream")}
        
        async with httpx.AsyncClient() as client:
            try:
                # 1. Anonymize
                response = await client.post(url_anon, files=files, timeout=60.0)
                if response.status_code == 200:
                    session_id = response.headers.get("X-Session-ID")
                    print(f"[ANON-SUCCESS] {filename} anonymized. Session: {session_id}")
                    out_path = f"test_results/anonymized_{filename}"
                    with open(out_path, "wb") as out_f:
                        out_f.write(response.content)
                        
                    # 2. Deanonymize
                    with open(out_path, "rb") as out_f:
                        deanon_files = {"file": (f"anonymized_{filename}", out_f, "application/octet-stream")}
                        deanon_data = {"session_id": session_id}
                        
                        resp_deanon = await client.post(url_deanon, files=deanon_files, data=deanon_data, timeout=60.0)
                        
                        if resp_deanon.status_code == 200:
                            print(f"[DEANON-SUCCESS] {filename} restored.")
                            deanon_path = f"test_results/restored_{filename}"
                            with open(deanon_path, "wb") as res_f:
                                res_f.write(resp_deanon.content)
                        else:
                            print(f"[DEANON-ERROR] {filename} failed with status {resp_deanon.status_code}: {resp_deanon.text}")
                else:
                    print(f"[ANON-ERROR] {filename} failed with status {response.status_code}: {response.text}")
            except Exception as e:
                print(f"[EXCEPTION] Failed to process {filename}: {e}")

async def main():
    print("Starting deep tests...")
    await asyncio.gather(
        test_file("test.txt"),
        test_file("test.docx"),
        test_file("test.xlsx"),
        test_file("test.pdf")
    )
    print("Deep tests completed.")

if __name__ == "__main__":
    asyncio.run(main())
