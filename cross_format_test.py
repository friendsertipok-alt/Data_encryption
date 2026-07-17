import os
import httpx
import asyncio
from docx import Document

os.makedirs("test_files", exist_ok=True)
os.makedirs("test_results", exist_ok=True)

async def test_cross_format():
    url_anon = "http://127.0.0.1:8001/api/anonymize/file"
    url_deanon = "http://127.0.0.1:8001/api/deanonymize/file"
    
    # Шаг 1: Создаем оригинальный TXT файл
    txt_filename = "original.txt"
    original_text = "Секретный агент: Джеймс Бонд. Почта: 007@mi6.uk"
    with open(f"test_files/{txt_filename}", "w", encoding="utf-8") as f:
        f.write(original_text)
        
    print("[1] Отправляем TXT на анонимизацию...")
    
    async with httpx.AsyncClient() as client:
        # Анонимизируем TXT
        with open(f"test_files/{txt_filename}", "rb") as f:
            files = {"file": (txt_filename, f, "text/plain")}
            response = await client.post(url_anon, files=files, timeout=60.0)
            
            if response.status_code == 200:
                session_id = response.headers.get("X-Session-ID")
                anonymized_text = response.text
                print(f"[SUCCESS] TXT анонимизирован. Session ID: {session_id}")
                print(f"[ANON] Анонимизированный текст: {anonymized_text}")
                
                # Шаг 2: Берем анонимизированный текст и вставляем его в новый DOCX
                docx_filename = "fake_document.docx"
                doc = Document()
                doc.add_paragraph("Это документ, созданный из текста TXT файла.")
                doc.add_paragraph(anonymized_text)
                doc.save(f"test_files/{docx_filename}")
                
                print("\n[2] Отправляем DOCX на деанонимизацию с тем же Session ID...")
                
                # Деанонимизируем DOCX
                with open(f"test_files/{docx_filename}", "rb") as doc_f:
                    deanon_files = {"file": (docx_filename, doc_f, "application/octet-stream")}
                    deanon_data = {"session_id": session_id}
                    
                    resp_deanon = await client.post(url_deanon, files=deanon_files, data=deanon_data, timeout=60.0)
                    
                    if resp_deanon.status_code == 200:
                        deanon_path = f"test_results/restored_{docx_filename}"
                        with open(deanon_path, "wb") as res_f:
                            res_f.write(resp_deanon.content)
                        print(f"[SUCCESS] DOCX успешно расшифрован и сохранен в {deanon_path}")
                        
                        # Читаем результат из DOCX, чтобы убедиться, что Джеймс Бонд вернулся
                        doc_restored = Document(deanon_path)
                        restored_text_full = "\n".join([p.text for p in doc_restored.paragraphs])
                        print(f"[RESTORED] Текст внутри расшифрованного DOCX:\n{restored_text_full}")
                    else:
                        print(f"[ERROR] Ошибка деанонимизации: {resp_deanon.text}")
            else:
                print(f"[ERROR] Ошибка анонимизации: {response.text}")

if __name__ == "__main__":
    asyncio.run(test_cross_format())
