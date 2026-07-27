import os
import httpx
import asyncio
from docx import Document
from openpyxl import Workbook
import random
import time

os.makedirs("stress_files", exist_ok=True)
os.makedirs("stress_results", exist_ok=True)

# Геренация сложного DOCX отчета
def generate_complex_docx():
    doc = Document()
    doc.add_heading('Аналитический Отчет по Безопасности №45-B', 0)
    
    doc.add_heading('1. Участники инцидента', level=1)
    doc.add_paragraph("Инцидент был инициирован Ивановым Иваном Ивановичем. "
                      "Ответственный за проверку: Петрову Петру Петровичу. "
                      "Копия направлена Сидоровым Сидором Сидоровичем.")
    
    doc.add_heading('2. Контакты и реквизиты', level=1)
    doc.add_paragraph("Email для связи: security_audit@megabank.ru. Телефон горячей линии: +7 (999) 123-45-67. "
                      "Подозрительные транзакции зафиксированы по карте 4276 1111 2222 3333, ИНН организации 7707083893.")
    
    doc.add_heading('3. Зарубежные партнеры', level=1)
    doc.add_paragraph("Also involved in the audit: Mr. John Doe (john.doe@fbi.gov). "
                      "Report provided by Alice Smith.")
                      
    doc.save("stress_files/complex_report.docx")
    print("[+] Сгенерирован complex_report.docx")

# Генерация огромной базы XLSX
def generate_huge_xlsx(rows=2000):
    wb = Workbook()
    ws = wb.active
    ws.title = "Client Database"
    
    headers = ["ID", "Full Name", "Email", "Phone", "INN", "Credit Card", "Notes"]
    ws.append(headers)
    
    first_names = ["Алексей", "Борис", "Владимир", "Григорий", "Дмитрий", "Евгений", "Жанна", "Зинаида"]
    last_names = ["Смирнов", "Иванов", "Кузнецов", "Попов", "Соколов", "Лебедев", "Козлов", "Новиков"]
    patronymics = ["Александрович", "Борисович", "Владимирович", "Иванович", "Петрович", "Сергеевич"]
    
    for i in range(1, rows + 1):
        fn = random.choice(first_names)
        ln = random.choice(last_names)
        pn = random.choice(patronymics)
        full_name = f"{ln} {fn} {pn}"
        
        email = f"user{i}_{random.randint(100,999)}@mail.ru"
        phone = f"+7 (9{random.randint(10,99)}) {random.randint(100,999)}-{random.randint(10,99)}-{random.randint(10,99)}"
        inn = f"77{random.randint(10000000, 99999999)}"
        card = f"4276 {random.randint(1000,9999)} {random.randint(1000,9999)} {random.randint(1000,9999)}"
        notes = f"Передал документы {full_name}у. Связаться по {email}."
        
        ws.append([i, full_name, email, phone, inn, card, notes])
        
    wb.save("stress_files/large_database.xlsx")
    print(f"[+] Сгенерирован large_database.xlsx на {rows} строк")

async def process_file(filename, mode, client):
    url_anon = "http://127.0.0.1:8001/api/anonymize/file"
    url_deanon = "http://127.0.0.1:8001/api/deanonymize/file"
    headers = {"X-API-Key": "test_key_123"}
    
    file_path = f"stress_files/{filename}"
    file_ext = filename.split('.')[-1]
    
    print(f"\n[{mode.upper()}] Начинаем обработку {filename}...")
    start_time = time.time()
    
    try:
        # АНОНИМИЗАЦИЯ
        with open(file_path, "rb") as f:
            files = {"file": (filename, f, "application/octet-stream")}
            data = {"mode": mode}
            # Увеличиваем таймаут для огромных файлов
            response = await client.post(url_anon, headers=headers, files=files, data=data, timeout=300.0)
            
            if response.status_code != 200:
                print(f"[ERROR] Анонимизация упала {filename}: {response.text}")
                return
                
            session_id = response.headers.get("X-Session-ID")
            anon_time = time.time() - start_time
            print(f"  [SUCCESS] Анонимизация успешна за {anon_time:.2f} сек. Session ID: {session_id}")
            
            out_path = f"stress_results/anon_{mode}_{filename}"
            with open(out_path, "wb") as out_f:
                out_f.write(response.content)
                
        # ДЕАНОНИМИЗАЦИЯ
        deanon_start = time.time()
        with open(out_path, "rb") as out_f:
            files = {"file": (f"anon_{mode}_{filename}", out_f, "application/octet-stream")}
            data = {"session_id": session_id}
            resp_deanon = await client.post(url_deanon, headers=headers, files=files, data=data, timeout=300.0)
            
            if resp_deanon.status_code != 200:
                print(f"[ERROR] Деанонимизация упала {filename}: {resp_deanon.text}")
                return
                
            deanon_time = time.time() - deanon_start
            deanon_path = f"stress_results/restored_{mode}_{filename}"
            with open(deanon_path, "wb") as res_f:
                res_f.write(resp_deanon.content)
                
            print(f"  [SUCCESS] Деанонимизация успешна за {deanon_time:.2f} сек. Файл сохранен.")
            
    except Exception as e:
        print(f"[EXCEPTION] {filename} в режиме {mode} упал с ошибкой: {str(e)}")

async def run_stress_test():
    print("=== ГЕНЕРАЦИЯ ДАННЫХ ДЛЯ СТРЕСС ТЕСТА ===")
    generate_complex_docx()
    # 200 строк = 200 x 6 ячеек = 1200 кусков текста. Вполне достаточно для стресс-теста.
    generate_huge_xlsx(200)
    
    print("\n=== ЗАПУСК СТРЕСС ТЕСТОВ ===")
    files_to_test = ["complex_report.docx", "large_database.xlsx"]
    modes = ["fake", "tags"]
    
    async with httpx.AsyncClient() as client:
        # Тестируем поочередно, чтобы не убить сервер ООМ (Out Of Memory)
        for mode in modes:
            for f in files_to_test:
                await process_file(f, mode, client)
                
    print("\n=== СТРЕСС ТЕСТ ЗАВЕРШЕН ===")

if __name__ == "__main__":
    asyncio.run(run_stress_test())
